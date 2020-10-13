import docx
import re
from docx.shared import Pt


document = docx
opa = docx.Document('C:/Users/kir/Desktop/dox/data.docx').paragraphs
opa_1 = docx.Document('C:/Users/kir/Desktop/dox/data_with_spaces.docx').paragraphs
keywords = {'ИНН': re.compile('ИНН?\s?[-:]?\s?(\d+)'),
            'СНИЛС':re.compile('СНИЛС?\s?[-:]?\s?(\d+)'),
            # 'Паспорт':re.compile('[Пп]аспорт?а??\s?[-:]?\s?(\d+)?\s(\d+)?'),
            'Паспорт':re.compile('[Пп]аспорт?а??\s?[-:]?\s?(\d+)'),
            'Фамилия':re.compile('[Фф]амилия?и??\s?[-:]?\s?(\w+)'),
            'Имя':re.compile('[Ии]мя?а??\s?[-:]?\s?(\w+)'),
            'Отчество':re.compile('[Оо]тчество?а??\s?[-:]?\s?(\w+)')
            }

new_info = {}
for i in opa:
    for k in keywords.keys():
        result = re.findall(keywords[k],i.text)
        if result:
            new_info[k] = result


print(new_info)
for i in new_info.keys():
    for s in opa_1:
        if i in s.text:
            s.text = re.sub(r'(' + i + '?\s?[-:]?\s?)_*',
                         r'\1 ' + new_info[i][0], s.text)



for l in opa_1:
    print(l.text)

new = docx.Document()
style = new.styles['Normal']
font = style.font
font.size = Pt(14)

for i in opa_1:
    new.add_paragraph(i.text)

new.save('C:/Users/kir/Desktop/dox/changed.docx')
